import os
import re
import sys
import time
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from concurrent.futures import ThreadPoolExecutor


def loading_animation():
    spinner = ['|', '/', '-', '\\']
    colors = ['\033[91m', '\033[92m', '\033[93m', '\033[94m']
    for _ in range(10):
        for j, char in enumerate(spinner):
            sys.stdout.write(f'\r{colors[j]}Loading... {char}\033[0m')
            sys.stdout.flush()
            time.sleep(0.1)


# HTML Parsing
def extract_data(file_location):
    try:
        with open(file_location, 'r', encoding='utf-8') as file:
            html_content = file.read()
    except FileNotFoundError:
        print(f"Error: File not found at location '{file_location}'")
        return

    soup = BeautifulSoup(html_content, 'html.parser')
    divs = soup.find_all('div')
    div_text = [div.text for div in divs]
    return div_text


# Text Processing (Vulnerability Extraction)
def extract_vulnerabilities(div_text):
    vulnerabilities = []
    for text in div_text:
        if "Div 1:" in text:
            for line in div_text:
                line = line.strip()
                if line == "Vulnerabilities by PluginExpand All | Collapse All":
                    break
                if line.startswith("- "):
                    vulnerability_name = line.split('-', 1)[-1].strip()
                    vulnerabilities.append(vulnerability_name)
            break  # Exit loop after finding vulnerabilities
    return vulnerabilities[4:]


def filter_vulnerabilities_by_risk(vulnerabilities, div_text):
    filtered_vulns = []
    for vuln in vulnerabilities:
        if any("Risk Factor" in line and vuln in line for line in div_text):
            filtered_vulns.append(vuln)
    return filtered_vulns


def select_vulnerabilities(vulnerabilities):
    print("List of vulnerabilities:")
    for vuln in vulnerabilities:
        print(vuln)

    selected = []
    while True:
        user_input = input("Enter vulnerability names to include (type 'exit' to stop): ")
        if user_input.lower() == 'exit':
            break
        if user_input in vulnerabilities:
            selected.append(user_input)
        else:
            print(f"Invalid vulnerability name: {user_input}")
    return selected


# Text Extraction Functions (Synopsis, Description, Solution, etc.)
def extract_text_after_second_occurrence(name, keyword, div_text):
    occurrences = "".join(div_text).split(name, 2)
    if len(occurrences) > 2:
        start_index = occurrences[2].find(keyword)
        if keyword == "Description":
            end_index = min(
                occurrences[2].find(x, start_index) for x in ["Solution", "See Also"] if x in occurrences[2])
        else:
            end_index = occurrences[2].find("\n\n", start_index)
        if start_index != -1 and end_index != -1:
            return occurrences[2][start_index:end_index].strip()
    return None


def extract_ip(name, div_text):
    solution = extract_text_after_second_occurrence(name, "Solution", div_text)
    if solution:
        return re.findall(r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\s+\([^)]+\)', solution)
    return None


def extract_cve(name, div_text):
    return extract_text_after_second_occurrence(name, "CVE", div_text)


# Word Document Creation
def set_cell_border(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['left', 'right', 'top', 'bottom']:
        element = OxmlElement(f'w:{border}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'auto')
        tcBorders.append(element)
    tcPr.append(tcBorders)


def create_docx_with_tables(vulnerabilities, div_text, file_name):
    doc = Document()
    for name in vulnerabilities:
        synopsis = extract_text_after_second_occurrence(name, "Synopsis", div_text)
        description = extract_text_after_second_occurrence(name, "Description", div_text)
        solution = extract_text_after_second_occurrence(name, "Solution", div_text)
        ips = extract_ip(name, div_text)
        risk_factor = extract_text_after_second_occurrence(name, "Risk Factor", div_text)
        cve = extract_cve(name, div_text)

        table = doc.add_table(rows=6, cols=3)
        table.autofit = False

        # Merge cells and set headers
        table.cell(0, 2).merge(table.cell(5, 2))
        table.cell(1, 0).merge(table.cell(1, 1))
        table.cell(2, 0).merge(table.cell(2, 1))
        table.cell(3, 0).merge(table.cell(3, 1))
        table.cell(4, 0).merge(table.cell(4, 1))
        table.cell(5, 0).merge(table.cell(5, 1))

        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell)

        table.cell(0, 0).text = "Vulnerability Name:\n" + name
        table.cell(1, 0).text = synopsis
        table.cell(2, 0).text = description
        table.cell(3, 0).text = " Proof of Concept:"
        table.cell(4, 0).text = solution
        affected_devices = "\n".join(ips) if ips else "N/A"
        table.cell(5, 0).text = f"Affected Device(s):\n{affected_devices}"
        table.cell(0, 1).text = risk_factor if risk_factor else "N/A"
        table.cell(0, 2).text = cve if cve else "N/A"

        if name != vulnerabilities[-1]:  # Add spacing between tables except for the last one
            doc.add_paragraph("\n\n")

    doc.save(f'Detail_{file_name}.docx')


def create_index_with_tables(vulnerabilities, div_text, file_name):
    doc = Document()

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'  # Apply a table style for better visual appeal
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Synopsis'
    hdr_cells[2].text = 'CVE'
    hdr_cells[3].text = 'Solution'
    hdr_cells[4].text = 'Affected Devices'
    hdr_cells[5].text = 'Risk Factor'

    for i, name in enumerate(vulnerabilities, start=1):
        row_cells = table.add_row().cells
        for cell in row_cells:
            set_cell_border(cell)
        row_cells[0].text = name
        row_cells[1].text = extract_text_after_second_occurrence(name, "Synopsis", div_text) or "N/A"
        row_cells[2].text = extract_cve(name, div_text) or "N/A"
        row_cells[3].text = extract_text_after_second_occurrence(name, "Solution", div_text) or "N/A"
        row_cells[4].text = "\n".join(extract_ip(name, div_text) or ["N/A"])
        row_cells[5].text = extract_text_after_second_occurrence(name, "Risk Factor", div_text) or "N/A"

    doc.save(f'Index_{file_name}.docx')
    # Main Execution
    if __name__ == "__main__":
        loading_animation()
        file_location = input("Enter the file location (e.g., 'C:/path/to/your/file.html'): ")
        div_text = extract_data(file_location)

        if div_text:
            file_name_with_extension = file_location.split('/')[-1]
            vulnerabilities = extract_vulnerabilities(div_text)
            filtered_vulns = filter_vulnerabilities_by_risk(vulnerabilities, div_text)
            selected_vulns = select_vulnerabilities(filtered_vulns)

            # Create Word documents
            with ThreadPoolExecutor() as executor:  # Using a thread pool for concurrency
                executor.submit(create_docx_with_tables, selected_vulns, div_text, file_name_with_extension)
                executor.submit(create_index_with_tables, selected_vulns, div_text, file_name_with_extension)

            # File cleanup logic remains the same (not repeated here)

        print("\nIt's Done! Enjoy your Word documents.")